import json
from anytree import Node, RenderTree
from docx import Document
from docx.shared import Pt

def build_tree_with_labels(node_data, parent=None, depth=0):
    """
    Builds an annotated anytree tree from AiZynthFinder JSON node.
    """
    node_type = node_data.get("type", "")
    smiles = node_data.get("smiles", "")
    in_stock = "[STOCK]" if node_data.get("in_stock", False) else ""
    
    if node_type == "reaction":
        label = f"[REACTION] {in_stock} {smiles}"
        metadata = node_data.get("metadata", {})
        template = metadata.get("template_code", "")
        policy_prob = metadata.get("policy_probability", "")
        label += f"\n          Template: {template}, Prob: {policy_prob}"
    else:
        label = f"[MOL] {in_stock} {smiles}"
    
    tree_node = Node(label, parent=parent, depth=depth)

    for child in node_data.get("children", []):
        build_tree_with_labels(child, tree_node, depth + 1)

    return tree_node

def export_to_word(json_file, output_docx):
    with open(json_file, "r") as f:
        data = json.load(f)

    doc = Document()
    doc.add_heading("Retrosynthesis Tree Output", level=1)

    for i, route in enumerate(data):
        route_score = route.get('scores', {}).get('state score', 'N/A')
        is_solved = route.get('metadata', {}).get('is_solved', False)
        doc.add_heading(f"Route {i+1}", level=2)
        doc.add_paragraph(f"Solved: {is_solved}, Score: {route_score}", style='Normal')

        root = build_tree_with_labels(route)
        for pre, _, node in RenderTree(root):
            indent_level = node.depth * 4  # 4 spaces per depth level
            text = f"{pre}{node.name}"
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(" " * indent_level + text)
            run.font.size = Pt(10)
    
    doc.save(output_docx)
    print(f"✅ Word file saved to '{output_docx}'")

# === Usage ===
export_to_word("output.json", "synthesis_tree.docx")
